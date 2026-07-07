from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("BimaHeadquarter_Business_Product_Report.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(25, 35, 45)
MUTED = RGBColor(92, 102, 112)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "E8EEF5"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D9DEE7", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(doc, text="", size=11, color=INK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size={1: 16, 2: 13, 3: 12}[level], color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        for run in p.runs:
            run.clear()
        r = p.add_run(item)
        set_font(r, size=10.8, color=INK)


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color="C9D3E1", size="8")
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_font(r, size=10.5, color=DARK_BLUE, bold=True)
    r = p.add_run(text)
    set_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    mark_header_row(table.rows[0])
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_FILL)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, size=10, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=9.7, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for level, size, color, before, after in [
        (1, 16, BLUE, 16, 8),
        (2, 13, BLUE, 12, 6),
        (3, 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = ""
    header.paragraph_format.space_after = Pt(0)
    hr = header.add_run("BimaHeadquarter CRM | Business & Product Report")
    set_font(hr, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("Powered by InsureDesk IMF Pvt. Ltd.")
    set_font(fr, size=8.5, color=MUTED)

    add_para(doc, "BimaHeadquarter CRM", size=25, color=DARK_BLUE, bold=True, after=3, before=24)
    add_para(doc, "Business & Product Report", size=15, color=MUTED, after=18)
    add_para(
        doc,
        "Prepared for business partners, clients, investors, and insurance professionals",
        size=10.5,
        color=INK,
        bold=True,
        after=2,
    )
    add_para(doc, "Developed by InsureDesk IMF Pvt. Ltd.", size=10.5, color=MUTED, after=18)
    add_callout(
        doc,
        "Positioning",
        "A cloud-based Insurance CRM that centralizes customers, policies, renewals, claims, endorsements, reporting, automation, and intelligent PDF processing for insurance businesses in India.",
    )
    add_para(doc, "Website: www.bimaheadquarter.com | Email: info@bimaheadquarter.com | Phone: +91 8818889660", size=9.5, color=MUTED, after=24)

    doc.add_page_break()

    add_heading(doc, "Executive Summary")
    add_para(
        doc,
        "BimaHeadquarter is a comprehensive cloud-based Insurance Customer Relationship Management platform designed to digitize and streamline insurance operations for agencies, brokers, POSPs, insurance consultants, and enterprises across India.",
    )
    add_para(
        doc,
        "The platform centralizes the full insurance lifecycle from customer acquisition and policy administration to renewals, claims, endorsements, reporting, and automation. By replacing fragmented spreadsheets and manual processes, it helps organizations improve operational efficiency, strengthen customer service, and support sustainable business growth.",
    )

    add_heading(doc, "Vision and Mission")
    add_para(doc, "Vision", size=11.5, color=DARK_BLUE, bold=True, after=3)
    add_para(doc, "To become India's most trusted digital operating system for insurance businesses by simplifying every aspect of insurance management through technology and automation.")
    add_para(doc, "Mission", size=11.5, color=DARK_BLUE, bold=True, after=3)
    add_bullets(
        doc,
        [
            "Digitize insurance operations and reduce manual workload.",
            "Improve customer experience and renewal retention.",
            "Simplify claim assistance and customer follow-up.",
            "Empower insurance professionals with intelligent technology.",
        ],
    )

    add_heading(doc, "Industry Challenges")
    add_para(doc, "Many insurance businesses still depend on traditional processes that create avoidable operational friction.")
    add_table(
        doc,
        ["Challenge Area", "Typical Issue", "Business Impact"],
        [
            ("Policy operations", "Manual tracking, Excel records, and paper-heavy documentation", "Higher turnaround time and inconsistent records"),
            ("Renewals", "Missed opportunities and fragmented follow-ups", "Lower retention and revenue leakage"),
            ("Claims", "Slow document collection and status tracking", "Poor customer visibility and service delays"),
            ("Management reporting", "Disconnected customer and policy data", "Limited insight into growth, productivity, and risk"),
        ],
        [2300, 3900, 3160],
    )

    add_heading(doc, "BimaHeadquarter Solution")
    add_para(
        doc,
        "BimaHeadquarter provides a unified digital platform that manages the complete insurance business workflow through a secure, scalable, and intelligent CRM. Customer records, policy information, renewals, claims, endorsements, documents, reports, analytics, and communication are consolidated into a single interface.",
    )
    add_callout(
        doc,
        "Operating model",
        "One system of record for customers, policies, teams, documents, follow-ups, and business performance.",
    )

    add_heading(doc, "Core Platform Modules")
    add_table(
        doc,
        ["Module", "Purpose", "Representative Capabilities"],
        [
            ("Dashboard", "Real-time business performance overview", "KPIs, graphical reports, notifications, recent activities, team performance"),
            ("Customer Management", "Centralized profiles and relationship history", "Contact details, policy history, documents, communication records, search and filters"),
            ("Policy Management", "Storage and tracking across insurance categories", "Premiums, coverage, expiry monitoring, documents, advanced search"),
            ("Renewal Management", "Retention-focused renewal workflow", "Upcoming and expired renewals, follow-up status, lost renewal analysis"),
            ("Claims Management", "Claim lifecycle coordination", "Registration, document collection, survey, processing, settlement tracking"),
            ("Endorsements", "Structured policy modification support", "Address changes, name corrections, financier updates, sum insured revisions, policy corrections"),
        ],
        [2100, 3000, 4260],
    )

    add_heading(doc, "Policy Lines Supported")
    add_bullets(
        doc,
        [
            "Motor, health, life, commercial, fire, marine, warehouse, engineering, and miscellaneous insurance.",
            "Policy storage, premium tracking, coverage management, expiry monitoring, document handling, and advanced search.",
        ],
    )

    add_heading(doc, "Intelligent PDF Processing")
    add_para(
        doc,
        "One of the platform's key differentiators is intelligent document processing. The system uses OCR and document intelligence to extract important information from insurance policy PDFs and presents the output for user review before saving.",
    )
    add_table(
        doc,
        ["Extracted Data", "Operational Value"],
        [
            ("Customer name, policy number, insurer, and policy dates", "Reduces repetitive entry and improves record consistency"),
            ("Premium, sum insured, coverage, and vehicle details", "Speeds policy creation and improves downstream reporting"),
            ("Review-before-save workflow", "Maintains human control while minimizing manual effort"),
        ],
        [3500, 5860],
    )

    add_heading(doc, "Manual Entry, Renewals, Claims, and Endorsements")
    add_para(
        doc,
        "For policies requiring manual processing, the CRM provides structured forms with validation, dynamic field loading, customer linking, and streamlined workflows. Renewal, claim, and endorsement modules then carry that data into recurring service processes.",
    )
    add_bullets(
        doc,
        [
            "Renewal workflow: upcoming renewals, expired policies, contact-wise and policy-wise tracking, renewal status, follow-up management, and lost renewal analysis.",
            "Claim workflow: registration, document collection, survey, processing, and settlement monitoring.",
            "Endorsement workflow: address changes, name corrections, financier updates, sum insured revisions, vehicle detail changes, and printable endorsement documents.",
        ],
    )

    add_heading(doc, "Reporting, Analytics, and Automation")
    add_para(
        doc,
        "The reporting layer helps management make informed decisions, while automation improves customer engagement and reduces repetitive follow-up work.",
    )
    add_table(
        doc,
        ["Capability", "Examples"],
        [
            ("Reports", "Sales, customer, policy, renewal, claim, team performance, premium analysis, and growth reports"),
            ("Dashboards", "Executive dashboards, KPIs, revenue views, renewals, claims, endorsements, and team performance"),
            ("Automation", "Renewal reminders, birthday greetings, follow-up notifications, customer updates, bulk messaging, and templates"),
        ],
        [2500, 6860],
    )

    add_heading(doc, "User Management and Security")
    add_para(
        doc,
        "The CRM includes role-based access control and secure authentication so organizations can manage access based on operational responsibility.",
    )
    add_table(
        doc,
        ["User Roles", "Security Features"],
        [
            ("Super Admin, Admin, Manager, Agent, Viewer", "Secure authentication, role-based permissions, data protection, audit logging, backup support, and cloud infrastructure"),
        ],
        [3900, 5460],
    )

    add_heading(doc, "Technology Stack")
    add_table(
        doc,
        ["Layer", "Technology"],
        [
            ("Frontend", "Next.js and React"),
            ("Backend", "Node.js"),
            ("Database and ORM", "PostgreSQL and Prisma"),
            ("Document Processing", "PDF processing and OCR technologies"),
            ("Cloud and Hosting", "Vercel and Google Drive integration"),
            ("Authentication", "JWT and Google Authentication"),
        ],
        [2800, 6560],
    )

    add_heading(doc, "Business Benefits")
    add_bullets(
        doc,
        [
            "Faster policy processing with reduced paperwork and fewer fragmented records.",
            "Improved renewal retention through structured tracking and reminders.",
            "Better customer management with centralized profiles, documents, and policy history.",
            "More efficient claims handling, endorsements, reporting, and team visibility.",
            "Scalable digital infrastructure that supports growth, compliance, and customer experience.",
        ],
    )

    add_heading(doc, "Future Roadmap")
    add_table(
        doc,
        ["Roadmap Theme", "Planned Enhancements"],
        [
            ("Customer and agent mobility", "Mobile application and customer self-service portal"),
            ("AI-assisted operations", "AI assistant, advanced OCR intelligence, AI business analytics, predictive renewal intelligence, conversational AI support"),
            ("Digital transaction layer", "Payment gateway integration and digital signatures"),
            ("Ecosystem connectivity", "Third-party API integrations and voice commands"),
        ],
        [3000, 6360],
    )

    add_heading(doc, "Conclusion")
    add_para(
        doc,
        "BimaHeadquarter is designed to modernize insurance operations by combining customer management, policy administration, automation, reporting, and intelligent document processing into a single platform. Its scalable architecture and operational focus make it a strong solution for insurance agencies and organizations seeking better productivity, customer satisfaction, and long-term business growth.",
    )
    add_para(
        doc,
        "By simplifying complex insurance workflows and enabling data-driven decision-making, BimaHeadquarter helps insurance professionals spend less time on administration and more time delivering value to customers.",
    )

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
